from openai import OpenAI
import PyPDF2
import re
import os
from docx import Document
from docx.shared import Pt

# variavel de ambiente contendo sua chave
client = OpenAI(
  api_key=os.environ['OPENAI_API_KEY'],  # chave da API
)


def extrair_dados_pdf(file_path):
    """Extrai texto do PDF fornecido."""
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        texto = ""
        for page in reader.pages:
            texto += page.extract_text()
    return texto


# pega o nome do cliente baseado no nome do arquivo_
def obter_cliente_do_nome(file_path):
    """Extrai o nome do cliente do nome do arquivo."""
    # Exemplo: CLIENTE_-_Assessment_de_Postura.05-09-2024-0300.pdf
    match = re.search(r"^([\w\-]+)", os.path.basename(file_path))
    if match:
        return match.group(1)
    return "Cliente Desconhecido"


def gerar_relatorio_executivo(cliente, texto, max_tokens=200000):
    """Gera um relatório executivo usando a nova API do ChatGPT."""
    mensagens = [
        {"role": "system", "content": "Você é um analista de segurança especializado em análise de vulnerabilidade"},
        {"role": "user", "content": f"""
        crie um relatório detalhado de análise de vulnerabilidades encontradas com as seguintes seções:


        “Descrição” descrever a vulnerabilidade e informar qual a plataforma e aplicativo afetado.
        "Requisitos para Exploração" informando o tipo de acesso necessário para explorar a vulnerabilidade e se é necessário ação do usuário e se para o sucesso é necessário que este usuário tenha permissões de administrador, nível de conhecimento técnico para exploração e se é necessário ferramenta específica e se está disponível publicamente;
        "Impacto Potencial" descrevendo o grau de comprometimento caso a ameaça tenha sucesso e qual a possível escalabilidade da ameaça;
        “Atualização do Software” descrevendo as ações necessárias para mitigar a vulnerabilidade e se existe algum plano de contingência caso a ação principal não possa ser feita de forma imediata. Informar também se já existe patch específico para correção da vulnerabilidade e se será necessário reiniciar o sistema.

        Lembre-se de se certificar da CVE no site do Mitre.
        
        segue um exemplo de relatório:


        Relatório de Análise de Vulnerabilidades:
        CVE-2023-29325:
        Descrição:
        A vulnerabilidade CVE-2023-29325 afeta o sistema operacional Windows 10 e 11. Essa vulnerabilidade permite que um atacante obtenha elevação de privilégios de usuário local para administrador, possibilitando o controle total do sistema.
        Avaliação do Sistema:
        Um atacante local com acesso de usuário padrão no sistema pode explorar essa vulnerabilidade. Não é necessária nenhuma ação do usuário, nem permissão de administrador para a exploração ter sucesso. O nível de conhecimento técnico requerido é médio, e existem ferramentas públicas disponíveis para exploração.
        Impacto Potencial:
        Caso a ameaça tenha sucesso, o atacante pode obter o controle total do sistema, com acesso de administrador. Isso permite que o invasor realize diversas ações maliciosas, como roubo de dados, instalação de malware, ou mesmo a tomada do controle total do computador.
        Atualização do Software:
        Para mitigar essa vulnerabilidade, é necessária a aplicação de um patch de segurança disponibilizado pela Microsoft. Esse patch deve ser instalado o mais rápido possível. Caso não seja possível a aplicação do patch de forma imediata, é recomendado isolar o sistema afetado da rede até que a correção possa ser aplicada. Reiniciar o sistema será necessário após a instalação do patch.
        tem menu de contexto
        {texto}
        """}
    ]
    #model="gpt-3.5-turbo"
    #model="gpt-4o-mini" vai ter que ser usado por motivos de rate limit
    #USE O MODELO "gpt-4o" SE VC TIVER tier 2+, O CONTEUDO DO RELATORIO VAI SER MELHOR
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=mensagens,
        max_tokens=max_tokens,
        n=1,
        stop=None,
        temperature=0.7,
    )


    return response.choices[0].message.content.strip()


def gerar_relatorio_para_docx(file_path, modelo_docx):
    """Gera um relatório executivo em um arquivo .docx usando um modelo fornecido."""
    texto_pdf = extrair_dados_pdf(file_path)
    cliente = obter_cliente_do_nome(file_path)
    relatorio_executivo = gerar_relatorio_executivo(cliente, texto_pdf)
   
    # Carrega o modelo fornecido
    doc = Document(modelo_docx)
   
    # Adiciona o relatório gerado ao documento
    doc.add_paragraph(f"Relatório Executivo - Cliente: {cliente}", style='Heading 1')
    p = doc.add_paragraph(relatorio_executivo)
    p.style.font.size = Pt(12)  # Define o tamanho da fonte para o texto adicionado
   
    # Salva o documento gerado com base no modelo
    output_path = f"{os.path.splitext(file_path)[0]}_Relatorio_Executivo.docx"
    doc.save(output_path)
   
    return output_path




def processar_multiplos_relatorios(diretorio_pdf, modelo_docx):
    """Processa vários arquivos PDF em um diretório e gera relatórios .docx para cada um."""
    for filename in os.listdir(diretorio_pdf):
        if filename.endswith(".pdf"):
            file_path = os.path.join(diretorio_pdf, filename)
            output_path = gerar_relatorio_para_docx(file_path, modelo_docx)
            print(f"Relatório gerado para {filename}: {output_path}")




# Caminho do diretório contendo os arquivos PDF
diretorio_pdf = "C:\\Users\\you\\Documents\\OPENAI\\Kasper"


# Caminho do modelo DOCX
modelo_docx = "modelo_relatorio.docx"  


# Processa todos os PDFs no diretório
processar_multiplos_relatorios(diretorio_pdf, modelo_docx)
